import os
import re
import hashlib
import csv
import logging
from typing import List, Tuple, Dict, Any, Optional
from datetime import datetime
import fitz  # PyMuPDF
from docx import Document
import pandas as pd
from .models import BusinessDocument, DocumentChunk, DocumentType, VersionStatus, ConfidentialityLevel

logger = logging.getLogger(__name__)

class DocumentProcessingError(Exception):
    """Raised when document processing fails"""
    pass

class DocumentProcessor:
    """Processes business documents and extracts structured information"""
    
    def __init__(self, chunk_size: int, chunk_overlap: int):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        logger.info(f"DocumentProcessor initialized - chunk_size: {chunk_size}, overlap: {chunk_overlap}")
    
    def read_pdf(self, file_path: str) -> Tuple[str, List[Tuple[str, int]]]:
        """Read content from PDF file and return text with page info"""
        try:
            doc = fitz.open(file_path)
            full_text = ""
            page_chunks = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                if page_text.strip():
                    full_text += page_text + "\n"
                    page_chunks.append((page_text, page_num + 1))
            
            doc.close()
            return full_text, page_chunks
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}", exc_info=True)
            raise DocumentProcessingError(f"Failed to read PDF file {file_path}: {str(e)}") from e
    
    def read_docx(self, file_path: str) -> str:
        """Read content from DOCX file"""
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}", exc_info=True)
            raise DocumentProcessingError(f"Failed to read DOCX file {file_path}: {str(e)}") from e
    
    def read_file(self, file_path: str) -> str:
        """
        Read content from text file (alias for read_text_file)

        Maintained for backward compatibility with tests.
        """
        return self.read_text_file(file_path)

    def read_text_file(self, file_path: str) -> str:
        """Read content from text file with UTF-8 encoding"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # If UTF-8 fails, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
    
    def read_excel(self, file_path: str) -> Tuple[str, List[Tuple[str, str, str]]]:
        """Read content from Excel file and return text with sheet/cell info"""
        try:
            # Load Excel file
            excel_file = pd.ExcelFile(file_path)
            full_text = ""
            sheet_chunks = []
            
            for sheet_name in excel_file.sheet_names:
                # Read each sheet
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert DataFrame to readable text format
                sheet_text = f"=== Sheet: {sheet_name} ===\n"
                
                # Add column headers
                headers = list(df.columns)
                sheet_text += "Columns: " + " | ".join(str(col) for col in headers) + "\n\n"
                
                # Add data rows (limit to first 1000 rows to avoid massive chunks)
                for idx, row in df.head(1000).iterrows():
                    row_text = " | ".join(str(value) if pd.notna(value) else "" for value in row.values)
                    sheet_text += f"Row {idx + 2}: {row_text}\n"  # +2 because Excel rows start at 1 and we have header
                
                # Determine cell range for this data
                max_row = min(len(df) + 1, 1001)  # +1 for header, max 1001
                max_col_letter = chr(ord('A') + len(df.columns) - 1)
                cell_range = f"A1:{max_col_letter}{max_row}"
                
                sheet_text += "\n"
                full_text += sheet_text
                sheet_chunks.append((sheet_text, sheet_name, cell_range))
            
            return full_text, sheet_chunks
        except Exception as e:
            raise Exception(f"Error reading Excel file {file_path}: {e}")
    
    def read_csv(self, file_path: str) -> Tuple[str, List[Tuple[str, int]]]:
        """Read content from CSV file and return text with row info"""
        try:
            # Read CSV file
            df = pd.read_csv(file_path)
            
            # Convert DataFrame to readable text format
            full_text = "=== CSV Data ===\n"
            
            # Add column headers
            headers = list(df.columns)
            full_text += "Columns: " + " | ".join(str(col) for col in headers) + "\n\n"
            
            # Track row chunks
            row_chunks = []
            
            # Add data rows (limit to first 1000 rows)
            for idx, row in df.head(1000).iterrows():
                row_text = " | ".join(str(value) if pd.notna(value) else "" for value in row.values)
                row_line = f"Row {idx + 2}: {row_text}\n"  # +2 because CSV rows start at 1 and we have header
                full_text += row_line
                row_chunks.append((row_line, idx + 2))
            
            return full_text, row_chunks
        except Exception as e:
            raise Exception(f"Error reading CSV file {file_path}: {e}")
    
    def calculate_content_hash(self, content: str) -> str:
        """Calculate SHA-256 hash of document content for deduplication"""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()
    


    def chunk_text(self, text: str) -> List[str]:
        """Split text into sentence-based chunks with overlap using config settings"""
        
        # Clean up the text
        text = re.sub(r'\s+', ' ', text.strip())  # Normalize whitespace
        
        # Better sentence splitting that handles abbreviations
        # This regex looks for periods followed by whitespace and capital letters
        # but ignores common abbreviations
        sentence_endings = re.compile(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\!|\?)\s+(?=[A-Z])')
        sentences = sentence_endings.split(text)
        
        # Clean sentences
        sentences = [s.strip() for s in sentences if s.strip()]
        
        chunks = []
        i = 0
        
        while i < len(sentences):
            current_chunk = []
            current_size = 0
            
            # Build chunk starting from sentence i
            for j in range(i, len(sentences)):
                sentence = sentences[j]
                
                # Calculate size with space
                space_size = 1 if current_chunk else 0
                total_addition = len(sentence) + space_size
                
                # Check if adding this sentence would exceed chunk size
                if current_size + total_addition > self.chunk_size and current_chunk:
                    break
                
                current_chunk.append(sentence)
                current_size += total_addition
            
            # Add chunk if we have content
            if current_chunk:
                chunks.append(' '.join(current_chunk))
                
                # Calculate overlap for next chunk
                if hasattr(self, 'chunk_overlap') and self.chunk_overlap > 0:
                    # Find how many sentences to overlap
                    overlap_size = 0
                    overlap_sentences = 0
                    
                    # Count backwards from end of current chunk
                    for k in range(len(current_chunk) - 1, -1, -1):
                        sentence_len = len(current_chunk[k]) + (1 if k < len(current_chunk) - 1 else 0)
                        if overlap_size + sentence_len <= self.chunk_overlap:
                            overlap_size += sentence_len
                            overlap_sentences += 1
                        else:
                            break
                    
                    # Move start position considering overlap
                    next_start = i + len(current_chunk) - overlap_sentences
                    i = max(next_start, i + 1)  # Ensure we make progress
                else:
                    # No overlap - move to next sentence after current chunk
                    i += len(current_chunk)
            else:
                # No sentences fit, move to next
                i += 1
        
        return chunks




    
    def process_business_document(self, file_path: str, doc_metadata: Dict[str, Any] = None) -> Tuple[BusinessDocument, List[DocumentChunk]]:
        """
        Process a business document (PDF, DOCX, TXT, XLSX, CSV) and extract content.
        Enhanced with metadata support and structured data handling.
        """
        filename = os.path.basename(file_path)
        file_extension = os.path.splitext(filename)[1].lower()
        
        # Determine file type and read content
        structured_chunks = []  # For Excel/CSV with specific location info
        
        if file_extension == '.pdf':
            file_type = 'pdf'
            content, page_chunks = self.read_pdf(file_path)
            # Convert page_chunks to standard format
            structured_chunks = [(chunk_text, {'page_number': page_num}) for chunk_text, page_num in page_chunks]
        elif file_extension == '.docx':
            file_type = 'docx'
            content = self.read_docx(file_path)
            structured_chunks = [(content, {})]
        elif file_extension in ['.txt', '.md']:
            file_type = 'txt'
            content = self.read_text_file(file_path)
            structured_chunks = [(content, {})]
        elif file_extension in ['.xlsx', '.xls']:
            file_type = 'xlsx'
            content, sheet_chunks = self.read_excel(file_path)
            # Convert sheet_chunks to standard format
            structured_chunks = [(chunk_text, {'sheet_name': sheet_name, 'cell_range': cell_range}) 
                                for chunk_text, sheet_name, cell_range in sheet_chunks]
        elif file_extension == '.csv':
            file_type = 'csv'
            content, row_chunks = self.read_csv(file_path)
            # Convert row_chunks to standard format
            structured_chunks = [(chunk_text, {'row_number': row_num}) for chunk_text, row_num in row_chunks]
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        if not content.strip():
            raise ValueError(f"No content extracted from {filename}")
        
        # Calculate content hash for deduplication
        content_hash = self.calculate_content_hash(content)
        
        # Get file metadata
        try:
            stat_result = os.stat(file_path)
            file_size = stat_result.st_size
            created_date = datetime.fromtimestamp(stat_result.st_ctime)
        except:
            file_size = None
            created_date = None
        
        # Use provided metadata or defaults
        if doc_metadata is None:
            doc_metadata = {}
        
        # Create business document object with enhanced metadata
        # Use try/except for enum conversions with fallback to defaults
        try:
            doc_type = DocumentType(doc_metadata.get('doc_type', 'index'))
        except ValueError as e:
            logger.warning(f"Invalid doc_type '{doc_metadata.get('doc_type')}' for {filename}, defaulting to 'index': {e}")
            doc_type = DocumentType.INDEX

        try:
            version = VersionStatus(doc_metadata.get('version', 'final'))
        except ValueError as e:
            logger.warning(f"Invalid version '{doc_metadata.get('version')}' for {filename}, defaulting to 'final': {e}")
            version = VersionStatus.FINAL

        try:
            confidentiality = ConfidentialityLevel(doc_metadata.get('confidentiality', 'internal'))
        except ValueError as e:
            logger.warning(f"Invalid confidentiality '{doc_metadata.get('confidentiality')}' for {filename}, defaulting to 'internal': {e}")
            confidentiality = ConfidentialityLevel.INTERNAL

        document = BusinessDocument(
            title=doc_metadata.get('title', filename),
            file_path=file_path,
            file_type=file_type,
            created_date=created_date,
            file_size=file_size,
            content_hash=content_hash,

            # Business metadata (from CSV or defaults with error handling)
            doc_type=doc_type,
            effective_date=doc_metadata.get('effective_date'),
            entities=doc_metadata.get('entities', []),
            address=doc_metadata.get('address'),
            version=version,
            confidentiality=confidentiality,
            notes=doc_metadata.get('notes'),

            # Legacy metadata field
            metadata=doc_metadata.get('metadata', {})
        )
        
        # Create document chunks with enhanced metadata
        document_chunks = []
        chunk_counter = 0
        
        # Process structured chunks based on file type
        for chunk_text, location_info in structured_chunks:
            if chunk_text.strip():
                # Use table-aware chunking for Excel/CSV, regular chunking for others
                if file_type in ['xlsx', 'csv']:
                    # For structured data, preserve the format and don't re-chunk
                    chunks = [chunk_text]
                else:
                    # Regular text chunking for PDFs, DOCX, TXT
                    chunks = self.chunk_text(chunk_text)
                
                for chunk in chunks:
                    # Enhanced source attribution
                    source_info = f"From {document.title}"
                    if 'page_number' in location_info:
                        source_info += f" (Page {location_info['page_number']})"
                    elif 'sheet_name' in location_info:
                        source_info += f" (Sheet: {location_info['sheet_name']}"
                        if 'cell_range' in location_info:
                            source_info += f", Range: {location_info['cell_range']}"
                        source_info += ")"
                    elif 'row_number' in location_info:
                        source_info += f" (Row {location_info['row_number']})"
                    
                    document_chunk = DocumentChunk(
                        content=f"{source_info}: {chunk}",
                        document_title=document.title,
                        file_type=file_type,
                        doc_type=document.doc_type,
                        chunk_index=chunk_counter,
                        
                        # Enhanced source attribution
                        page_number=location_info.get('page_number'),
                        sheet_name=location_info.get('sheet_name'),
                        cell_range=location_info.get('cell_range'),
                        row_number=location_info.get('row_number'),
                        
                        # Business metadata for filtering
                        effective_date=document.effective_date,
                        entities=document.entities,
                        version=document.version
                    )
                    document_chunks.append(document_chunk)
                    chunk_counter += 1
        
        return document, document_chunks
